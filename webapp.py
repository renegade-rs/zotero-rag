#!/usr/bin/env python3
"""Web interface for Zotero RAG with AI-powered Q&A.

Supports multiple LLM providers: Anthropic Claude, OpenAI, or Ollama (free/local).

Usage:
    source .venv/bin/activate
    python webapp.py

Then open http://localhost:5001 in your browser.
"""

import json
import logging
import os
import re
import subprocess
import sys
from datetime import datetime
from io import BytesIO

# ODF/OOo imports for file export
from odf.opendocument import OpenDocumentText
from odf.style import Style, TextProperties, ParagraphProperties
from odf.text import H, P, List, ListItem, Span, A
from odf.table import Table, TableRow, TableCell, TableColumn, TableColumns
from odf.namespaces import FONS, TEXTNS

# DOCX imports
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_UNDERLINE, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from dotenv import load_dotenv
load_dotenv()

from fastapi import FastAPI, Request, HTTPException, Depends
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse, HTMLResponse, Response
from fastapi.security import OAuth2PasswordRequestForm
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from auth import (
    authenticate_user, create_access_token, get_current_user,
    create_user, get_user as get_auth_user, get_all_users,
    approve_user, delete_user
)
from chat_history import (
    get_user_chats, get_chat, create_chat, add_message,
    update_chat_title, save_sources, delete_chat, sync_chat,
    get_chat_for_download
)
from src.config import (
    LLM_PROVIDER, ANTHROPIC_API_KEY, ANTHROPIC_MODEL,
    OPENAI_API_KEY, OPENAI_CHAT_MODEL,
    OLLAMA_BASE_URL, OLLAMA_CHAT_MODEL,
    ARCHIVE_ALIASES_FILE, BASE_URL,
)
from src.search_pipeline import init_pipeline, run_search, get_archive_aliases
from src.logging_config import setup_logging

log_file = setup_logging("webapp")
print(f"Logs written to: {log_file}")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()

class RegisterRequest(BaseModel):
    email: str
    username: str
    password: str
# Mount static files directory to serve JavaScript and other static assets
app.mount("/static", StaticFiles(directory="static"), name="static")

SYSTEM_PROMPT = """You are a research assistant helping a scholar with their research. You answer questions using ONLY the provided source documents.

Rules:
1. Answer ONLY from the provided sources. Do not use outside knowledge.
2. Cite sources using [N] notation matching the source numbers provided.
3. If the sources are insufficient to fully answer the question, explicitly state what information is missing.
4. Be neutral and professional. Report what the sources say without editorializing.
5. When quoting, use exact text from the sources.
6. End your response with a "Sources" section listing all cited sources.

The user may ask follow-up questions -- use conversation context plus any new sources provided."""


def _build_source_context(sources):
    """Format search results into a context block for the LLM."""
    if not sources:
        return "\n[No sources found for this query.]\n"
    parts = ["\n--- BEGIN SOURCES ---"]
    for s in sources:
        source_num = s.get('source_num') or 1
        authors = ', '.join(s.get('authors', [])) or 'Unknown'
        title = s.get('title', 'Untitled')
        date = s.get('date', '')
        item_type = s.get('item_type', '')
        text = s.get('text', '')

        page_start = s.get('page_start', 0)
        page_end = s.get('page_end', 0)
        
        work_pages = s.get('pages', '')
        pdf_page_str = ''
        if page_start > 0:
            pdf_page_str = f" [PDF pp. {page_start}-{page_end}]" if page_end > page_start else f" [PDF p. {page_start}]"
        
        page_str = ''
        if work_pages:
            page_str = f", {work_pages}" + (pdf_page_str if pdf_page_str else "")
        elif page_start > 0:
            page_str = f", pp. {page_start}-{page_end}" if page_end > page_start else f", p. {page_start}"

        archive = s.get('archive', '')
        archive_loc = s.get('archive_location', '')
        archive_str = ''
        if archive:
            archive_str = f"\nArchive: {archive}"
            if archive_loc:
                archive_str += f", {archive_loc}"

        parts.append(
            f"\n[{source_num}] \"{title}\" -- {authors}"
            f"{' (' + date + ')' if date else ''}{page_str}"
            f"\nType: {item_type}{archive_str}"
            f"\n{text}\n"
        )
    parts.append("--- END SOURCES ---\n")
    return '\n'.join(parts)


def _format_source_for_client(result):
    """Convert a search result dict into the format sent to the frontend."""
    meta = result['metadata']
    attachment_key = meta.get('attachment_key', '')
    attachment_type = meta.get('attachment_type', 'pdf')
    zotero_key = meta.get('zotero_key', '')
    pdf_page = meta.get('pdf_page', meta.get('page_start', 0))

    zotero_url = ''
    if attachment_key:
        zotero_url = f"/zotero/pdf/{attachment_key}"
        if pdf_page and attachment_type == 'pdf':
            zotero_url += f"?page={pdf_page}"
    elif zotero_key:
        zotero_url = f"/zotero/item/{zotero_key}"

    return {
        'title': meta.get('title', 'Untitled'),
        'authors': meta.get('authors', []),
        'date': meta.get('date', ''),
        'item_type': meta.get('item_type', ''),
        'archive': meta.get('archive', ''),
        'archive_location': meta.get('archive_location', ''),
        'page_start': meta.get('page_start', 0),
        'page_end': meta.get('page_end', 0),
        'pages': meta.get('pages', ''),
        'text': meta.get('text', '')[:800],
        'zotero_url': zotero_url,
        'score': float(result.get('score', 0)),
        'rerank_score': float(result['rerank_score']) if result.get('rerank_score') is not None else None,
        'chunk_id': result.get('id'),
        'zotero_key': zotero_key,
        'chunk_index': meta.get('chunk_index', 0),
    }

def _parse_markdown_to_odt_content(markdown_text: str) -> str:
    """Convert markdown to HTML for ODT conversion.
    Strips most markdown syntax while preserving structure."""
    
    text = markdown_text
    
    # Convert headers (preserve hierarchy)
    text = re.sub(r'^######\s+(.+)$', '<h6>\\1</h6>', text, flags=re.MULTILINE)
    text = re.sub(r'^#####\s+(.+)$', '<h5>\\1</h5>', text, flags=re.MULTILINE)
    text = re.sub(r'^####\s+(.+)$', '<h4>\\1</h4>', text, flags=re.MULTILINE)
    text = re.sub(r'^###\s+(.+)$', '<h3>\\1</h3>', text, flags=re.MULTILINE)
    text = re.sub(r'^##\s+(.+)$', '<h2>\\1</h2>', text, flags=re.MULTILINE)
    text = re.sub(r'^#\s+(.+)$', '<h1>\\1</h1>', text, flags=re.MULTILINE)
    
    # Convert bold/italic
    text = re.sub(r'\*\*\*(.+?)\*\*\*', '<strong><em>\\1</em></strong>', text)
    text = re.sub(r'___(.+?)___', '<strong><em>\\1</em></strong>', text)
    text = re.sub(r'\*\*(.+?)\*\*', '<strong>\\1</strong>', text)
    text = re.sub(r'__(.+?)__', '<strong>\\1</strong>', text)
    text = re.sub(r'\*(.+?)\*', '<em>\\1</em>', text)
    text = re.sub(r'_(.+?)_', '<em>\\1</em>', text)
    
    # Convert code blocks
    text = re.sub(r'```(\w+)?\n(.+?)```', '<pre>\\2</pre>', text, flags=re.DOTALL)
    text = re.sub(r'`(.+?)`', '<code>\\1</code>', text)
    
    # Convert blockquotes
    text = re.sub(r'^>\s+(.+)$', '<blockquote>\\1</blockquote>', text, flags=re.MULTILINE)
    
    # Convert unordered lists
    text = re.sub(r'^\s{0,3}\*\s+(.+)$', '<ul><li>\\1</li></ul>', text, flags=re.MULTILINE)
    text = re.sub(r'^\s{0,3}-\s+(.+)$', '<ul><li>\\1</li></ul>', text, flags=re.MULTILINE)
    
    # Convert ordered lists
    text = re.sub(r'^\s{0,3}\d+\.\s+(.+)$', '<ol><li>\\1</li></ol>', text, flags=re.MULTILINE)
    
    # Convert links
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', '<a href="\\2">\\1</a>', text)
    
    # Convert horizontal rules
    text = re.sub(r'^---+$', '<hr/>', text, flags=re.MULTILINE)
    text = re.sub(r'^\*\*\*+$', '<hr/>', text, flags=re.MULTILINE)
    
    return text


def _format_sources_section(sources: list) -> str:
    """Format sources as markdown for the sources section."""
    if not sources:
        return ""
    
    lines = ["\n## Sources\n"]
    for src in sources:
        authors = ', '.join(src.get('authors', [])) or 'Unknown'
        title = src.get('title', 'Untitled')
        date = src.get('date', '')
        item_type = src.get('item_type', '')
        page_info = ''
        if src.get('page_start'):
            page_end = src.get('page_end', src.get('page_start'))
            if page_end > src.get('page_start', 0):
                page_info = f", pp. {src['page_start']}-{page_end}"
            else:
                page_info = f", p. {src['page_start']}"
        
        lines.append(f"### [{src.get('source_num', '')}] {title}")
        lines.append(f"- **Authors:** {authors}")
        if date:
            lines.append(f"- **Date:** {date}")
        lines.append(f"- **Type:** {item_type}{page_info}")
        if src.get('archive'):
            lines.append(f"- **Archive:** {src['archive']}")
        if src.get('archive_location'):
            lines.append(f"- **Location:** {src['archive_location']}")
        if src.get('text'):
            preview = src['text']
            lines.append(f"- **Preview:** {preview}")
        lines.append("")
    
    return '\n'.join(lines)


def _add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    run = OxmlElement('w:r')
    run.append(OxmlElement('w:rPr'))
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0066cc')
    run.find(qn('w:rPr')).append(color)
    
    r = OxmlElement('w:t')
    r.text = text
    run.append(r)
    
    hyperlink.append(run)
    paragraph.add_run().element.addprevious(hyperlink)
    
    return paragraph


def _sanitize_text_for_export(text: str) -> str:
    """Remove control characters that are not XML-compatible."""
    if not text:
        return text
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)


def _chat_to_odt(chat_data: dict) -> BytesIO:
    """Convert chat messages and sources to ODT format."""
    
    doc = OpenDocumentText()
    
    # Style definitions
    h1_style = Style(name="H1", family="paragraph")
    h1_style.addElement(TextProperties(fontsize='18pt', fontweight='bold'))
    doc.styles.addElement(h1_style)
    
    h2_style = Style(name="H2", family="paragraph")
    h2_style.addElement(TextProperties(fontsize='16pt', fontweight='bold'))
    doc.styles.addElement(h2_style)
    
    h3_style = Style(name="H3", family="paragraph")
    h3_style.addElement(TextProperties(fontsize='14pt', fontweight='bold'))
    doc.styles.addElement(h3_style)
    
    user_msg_style = Style(name="UserMessage", family="paragraph")
    user_msg_style.addElement(ParagraphProperties(marginleft='20pt'))
    user_msg_style.addElement(TextProperties(color='#1a1a1a', fontweight='bold'))
    doc.styles.addElement(user_msg_style)
    
    assistant_msg_style = Style(name="AssistantMessage", family="paragraph")
    assistant_msg_style.addElement(ParagraphProperties(marginleft='0pt'))
    doc.styles.addElement(assistant_msg_style)
    
    source_style = Style(name="Source", family="paragraph")
    source_style.addElement(ParagraphProperties(marginleft='20pt'))
    source_style.addElement(TextProperties(fontsize='9pt', color='#666'))
    doc.styles.addElement(source_style)
    
    list_unordered_style = Style(name="ListUnordered", family="paragraph")
    list_unordered_style.addElement(ParagraphProperties(marginleft='3.8cm'))
    doc.styles.addElement(list_unordered_style)
    
    list_ordered_style = Style(name="ListOrdered", family="paragraph")
    list_ordered_style.addElement(ParagraphProperties(marginleft='3.8cm'))
    doc.styles.addElement(list_ordered_style)
    
    listitem_style = Style(name="ListItem", family="paragraph")
    listitem_style.addElement(ParagraphProperties(marginleft='2.5cm'))
    doc.styles.addElement(listitem_style)
    
    bold_style = Style(name="bold", family="text")
    bold_style.addElement(TextProperties(fontweight='bold'))
    doc.styles.addElement(bold_style)
    
    italic_style = Style(name="italic", family="text")
    italic_style.addElement(TextProperties(fontstyle='italic'))
    doc.styles.addElement(italic_style)
    
    bold_italic_style = Style(name="bold-italic", family="text")
    bold_italic_style.addElement(TextProperties(fontweight='bold', fontstyle='italic'))
    doc.styles.addElement(bold_italic_style)
    
    code_style = Style(name="code", family="text")
    code_style.addElement(TextProperties(fontfamily='monospace', fontsize='10pt', color='#333'))
    doc.styles.addElement(code_style)
    
    preview_style = Style(name="Preview", family="text")
    preview_style.addElement(TextProperties(color='#0066cc', fontsize='8pt'))
    doc.styles.addElement(preview_style)
    
    # Build chat content
    title = chat_data.get('title', 'Chat')
    messages = chat_data.get('messages', [])
    sources = chat_data.get('sources', [])
    
    # Add title
    doc.text.addElement(H(outlinelevel=1, text=title))
    
    # Add date/time
    doc.text.addElement(P(text=f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"))
    doc.text.addElement(P(text=" "))
    
    # Add messages
    for msg in messages:
        role = msg.get('role', '')
        content = msg.get('content', '')
        
        if role == 'user':
            # User message
            p = P(stylename=user_msg_style)
            p.addText(f"User: {content}")
            doc.text.addElement(p)
        elif role == 'assistant':
            # Assistant message - parse markdown with proper ODT structures
            import re
            
            lines = content.split('\n')
            i = 0
            while i < len(lines):
                line = lines[i]
                
                # Headers
                if line.startswith('# '):
                    text = re.sub(r'^#\s+(.+)$', r'\1', line)
                    doc.text.addElement(H(outlinelevel=1, text=text))
                    i += 1
                    continue
                
                if line.startswith('## '):
                    text = re.sub(r'^##\s+(.+)$', r'\1', line)
                    doc.text.addElement(H(outlinelevel=2, text=text))
                    i += 1
                    continue
                
                if line.startswith('### '):
                    text = re.sub(r'^###\s+(.+)$', r'\1', line)
                    doc.text.addElement(H(outlinelevel=3, text=text))
                    i += 1
                    continue
                
                # Blockquotes
                if line.startswith('> '):
                    text = re.sub(r'^>\s+(.+)$', r'\1', line)
                    p = P(stylename=assistant_msg_style)
                    _parse_inline_markdown(p, text)
                    doc.text.addElement(p)
                    i += 1
                    continue
                
                # Unordered lists (bullets)
                if re.match(r'^\s*\*\s+', line) or re.match(r'^\s*-\s+', line):
                    mylist = List()
                    while i < len(lines) and (re.match(r'^\s*\*\s+', lines[i]) or re.match(r'^\s*-\s+', lines[i])):
                        list_item_text = re.sub(r'^\s*[\*-]\s+', '', lines[i])
                        item = ListItem()
                        p = P()
                        _parse_inline_markdown(p, list_item_text)
                        item.addElement(p)
                        mylist.addElement(item)
                        i += 1
                    doc.text.addElement(mylist)
                    continue
                
                # Ordered lists (numbers)
                if re.match(r'^\s*\d+\.\s+', line):
                    mylist = List()
                    while i < len(lines) and re.match(r'^\s*\d+\.\s+', lines[i]):
                        list_item_text = re.sub(r'^\s*\d+\.\s+', '', lines[i])
                        item = ListItem()
                        p = P()
                        _parse_inline_markdown(p, list_item_text)
                        item.addElement(p)
                        mylist.addElement(item)
                        i += 1
                    doc.text.addElement(mylist)
                    continue
                
                # Tables (| syntax)
                if re.match(r'^\s*\|.*\|\s*$', line):
                    table, row_count = _parse_markdown_table(lines[i:])
                    if table:
                        doc.text.addElement(table)
                        i += row_count
                    else:
                        p = P(stylename=assistant_msg_style)
                        p.addText(line)
                        doc.text.addElement(p)
                        i += 1
                    continue
                
                # Regular paragraphs
                if line.strip():
                    p = P(stylename=assistant_msg_style)
                    _parse_inline_markdown(p, line)
                    doc.text.addElement(p)
                
                i += 1
        
        doc.text.addElement(P(text=" "))
    
    # Add sources section if available
    if sources:
        doc.text.addElement(H(outlinelevel=2, text="Sources"))
        
        for i, src in enumerate(sources, 1):
            p = P(stylename=source_style)
            
            authors = ', '.join(src.get('authors', [])) or 'Unknown'
            title = src.get('title', 'Untitled')
            date = src.get('date', '')
            
            source_text = f"[{i}] {title} -- {authors}"
            if date:
                source_text += f" ({date})"
            
            p.addText(source_text)
            doc.text.addElement(p)
            
            if src.get('page_start'):
                p_page = P(text=f"  Pages: {src['page_start']}")
                if src.get('page_end') and src['page_end'] != src['page_start']:
                    p_page.addText(f"-{src['page_end']}")
                doc.text.addElement(p_page)
            
            if src.get('item_type'):
                p_type = P(text=f"  Type: {src['item_type']}")
                doc.text.addElement(p_type)
            
            if src.get('archive'):
                p_archive = P(text=f"  Archive: {src['archive']}")
                doc.text.addElement(p_archive)
            
            if src.get('text'):
                preview_text = _sanitize_text_for_export(src['text']).replace('\n\n', '\n')
                if preview_text:
                    p_preview = P(text="  Preview: ")
                    run_preview = Span(text=preview_text)
                    run_preview.setAttrNS(TEXTNS, 'style-name', 'Preview')
                    p_preview.addElement(run_preview)
                    doc.text.addElement(p_preview)
            
            if src.get('zotero_url'):
                zotero_path = src['zotero_url']
                if not zotero_path.startswith('http://') and not zotero_path.startswith('https://'):
                    full_zotero_url = f"{BASE_URL}{zotero_path}"
                else:
                    full_zotero_url = zotero_path
                p_zotero = P()
                prefix_span = Span(text="  Open in Zotero: ")
                p_zotero.addElement(prefix_span)
                link = A(href=full_zotero_url, text=full_zotero_url)
                p_zotero.addElement(link)
                doc.text.addElement(p_zotero)
            
            p_sep = P(text="  ")
            doc.text.addElement(p_sep)
    
    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output


def _parse_inline_markdown(paragraph: P, text: str) -> P:
    """Parse inline markdown (bold, italic, code) and add to paragraph as spans."""
    import re
    
    while text:
        match = None
        
        # Bold + italic (***/___)
        match = re.match(r'\*\*\*(.+?)\*\*\*', text)
        if match:
            span = Span(text=match.group(1))
            span.setAttrNS(TEXTNS, 'style-name', 'bold-italic')
            paragraph.addElement(span)
            text = text[len(match.group(0)):]
            continue
        
        match = re.match(r'___(.+?)___', text)
        if match:
            span = Span(text=match.group(1))
            span.setAttrNS(TEXTNS, 'style-name', 'bold-italic')
            paragraph.addElement(span)
            text = text[len(match.group(0)):]
            continue
        
        # Bold (** or __)
        match = re.match(r'\*\*(.+?)\*\*', text)
        if match:
            span = Span(text=match.group(1))
            span.setAttrNS(TEXTNS, 'style-name', 'bold')
            paragraph.addElement(span)
            text = text[len(match.group(0)):]
            continue
        
        match = re.match(r'__(.+?)__', text)
        if match:
            span = Span(text=match.group(1))
            span.setAttrNS(TEXTNS, 'style-name', 'bold')
            paragraph.addElement(span)
            text = text[len(match.group(0)):]
            continue
        
        # Italic (* or _)
        match = re.match(r'\*(.+?)\*', text)
        if match:
            span = Span(text=match.group(1))
            span.setAttrNS(TEXTNS, 'style-name', 'italic')
            paragraph.addElement(span)
            text = text[len(match.group(0)):]
            continue
        
        match = re.match(r'_(.+?)_', text)
        if match:
            span = Span(text=match.group(1))
            span.setAttrNS(TEXTNS, 'style-name', 'italic')
            paragraph.addElement(span)
            text = text[len(match.group(0)):]
            continue
        
        # Inline code
        match = re.match(r'`(.+?)`', text)
        if match:
            span = Span(text=match.group(1))
            span.setAttrNS(TEXTNS, 'style-name', 'code')
            paragraph.addElement(span)
            text = text[len(match.group(0)):]
            continue
        
        # Regular text
        match = re.match(r'([^*_\`]+)', text)
        if match:
            paragraph.addText(match.group(1))
            text = text[len(match.group(0)):]
            continue
    
    return paragraph


def _parse_markdown_table(lines: list) -> tuple:
    """Parse markdown table and return (Table element, number of lines consumed)."""
    import re
    
    if not lines or not re.match(r'^\s*\|.*\|\s*$', lines[0]):
        return (None, 0)
    
    # Parse header
    header_line = lines[0].strip()
    if header_line.endswith('|'):
        header_line = header_line[:-1]
    if header_line.startswith('|'):
        header_line = header_line[1:]
    headers = [h.strip() for h in header_line.split('|')]
    
    # Parse separator (optional)
    separator_lines = 1
    if len(lines) > 1 and re.match(r'^\s*\|\s*[-:]+', lines[1]):
        separator_lines = 2
    
    # Parse rows
    rows = []
    i = separator_lines
    while i < len(lines) and re.match(r'^\s*\|.*\|\s*$', lines[i]):
        row_line = lines[i].strip()
        if row_line.endswith('|'):
            row_line = row_line[:-1]
        if row_line.startswith('|'):
            row_line = row_line[1:]
        cells = [c.strip() for c in row_line.split('|')]
        rows.append(cells)
        i += 1
    
    # Create ODT table with explicit column definitions
    table = Table(name="Table1")
    
    # Define columns before adding rows (required for multi-column ODT tables)
    col_defs = TableColumns()
    for _ in headers:
        col_defs.addElement(TableColumn())
    table.addElement(col_defs)
    
    # Header row
    header_row = TableRow()
    for header in headers:
        cell = TableCell()
        p = P()
        _parse_inline_markdown(p, header)
        p.setAttrNS(FONS, 'margin-left', '0.1in')
        p.setAttrNS(FONS, 'margin-right', '0.1in')
        cell.addElement(p)
        header_row.addElement(cell)
    table.addElement(header_row)
    
    # Data rows
    for row in rows:
        data_row = TableRow()
        for cell_text in row:
            cell = TableCell()
            p = P()
            _parse_inline_markdown(p, cell_text)
            p.setAttrNS(FONS, 'margin-left', '0.1in')
            p.setAttrNS(FONS, 'margin-right', '0.1in')
            cell.addElement(p)
            data_row.addElement(cell)
        table.addElement(data_row)
    
    return (table, i)


def _parse_inline_markdown_docx(paragraph, text: str):
    """Parse inline markdown and add to docx paragraph as runs."""
    import re
    
    while text:
        match = None
        
        # Bold + italic (***/___)
        match = re.match(r'\*\*\*(.+?)\*\*\*', text)
        if match:
            run = paragraph.add_run(match.group(1))
            run.bold = True
            run.italic = True
            text = text[len(match.group(0)):]
            continue
        
        match = re.match(r'___(.+?)___', text)
        if match:
            run = paragraph.add_run(match.group(1))
            run.bold = True
            run.italic = True
            text = text[len(match.group(0)):]
            continue
        
        # Bold (** or __)
        match = re.match(r'\*\*(.+?)\*\*', text)
        if match:
            run = paragraph.add_run(match.group(1))
            run.bold = True
            text = text[len(match.group(0)):]
            continue
        
        match = re.match(r'__(.+?)__', text)
        if match:
            run = paragraph.add_run(match.group(1))
            run.bold = True
            text = text[len(match.group(0)):]
            continue
        
        # Italic (* or _)
        match = re.match(r'\*(.+?)\*', text)
        if match:
            run = paragraph.add_run(match.group(1))
            run.italic = True
            text = text[len(match.group(0)):]
            continue
        
        match = re.match(r'_(.+?)_', text)
        if match:
            run = paragraph.add_run(match.group(1))
            run.italic = True
            text = text[len(match.group(0)):]
            continue
        
        # Underline (_text_)
        match = re.match(r'_([^_\s]+)_', text)
        if match:
            run = paragraph.add_run(match.group(1))
            run.underline = WD_UNDERLINE.SINGLE
            text = text[len(match.group(0)):]
            continue
        
        # Inline code
        match = re.match(r'`(.+?)`', text)
        if match:
            run = paragraph.add_run(match.group(1))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            text = text[len(match.group(0)):]
            continue
        
        # Regular text
        match = re.match(r'([^*_\`]+)', text)
        if match:
            paragraph.add_run(match.group(1))
            text = text[len(match.group(0)):]
            continue
    
    return paragraph


def _parse_markdown_table_docx(document, lines: list):
    """Parse markdown table and return docx table."""
    import re
    
    if not lines or not re.match(r'^\s*\|.*\|\s*$', lines[0]):
        return None
    
    # Parse header
    header_line = lines[0].strip()
    if header_line.endswith('|'):
        header_line = header_line[:-1]
    if header_line.startswith('|'):
        header_line = header_line[1:]
    headers = [h.strip() for h in header_line.split('|')]
    
    # Parse separator (optional)
    separator_lines = 1
    if len(lines) > 1 and re.match(r'^\s*\|\s*[-:]+', lines[1]):
        separator_lines = 2
    
    # Parse rows
    rows = []
    i = separator_lines
    while i < len(lines) and re.match(r'^\s*\|.*\|\s*$', lines[i]):
        row_line = lines[i].strip()
        if row_line.endswith('|'):
            row_line = row_line[:-1]
        if row_line.startswith('|'):
            row_line = row_line[1:]
        cells = [c.strip() for c in row_line.split('|')]
        rows.append(cells)
        i += 1
    
    # Create docx table
    table = document.add_table(rows=0, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    header_row = table.add_row()
    for j, header in enumerate(headers):
        cell = header_row.cells[j]
        p = cell.paragraphs[0]
        _parse_inline_markdown_docx(p, header)
    
    # Data rows
    for row in rows:
        data_row = table.add_row()
        for j, cell_text in enumerate(row):
            cell = data_row.cells[j]
            p = cell.paragraphs[0]
            _parse_inline_markdown_docx(p, cell_text)
    
    return (table, i)


def _chat_to_docx(chat_data: dict) -> BytesIO:
    """Convert chat messages and sources to DOCX format."""
    
    doc = Document()
    
    # Add styles
    styles = doc.styles
    
    # Heading 1 style (for main title)
    if 'Heading 1' not in styles:
        h1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.base_style = styles['Normal']
        h1_font = h1_style.font
        h1_font.size = Pt(18)
        h1_font.bold = True
    
    # Heading 2 style (for H2)
    if 'Heading 2' not in styles:
        h2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.base_style = styles['Normal']
        h2_font = h2_style.font
        h2_font.size = Pt(14)
        h2_font.bold = True
    
    # Heading 3 style (for H3)
    if 'Heading 3' not in styles:
        h3_style = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        h3_style.base_style = styles['Normal']
        h3_font = h3_style.font
        h3_font.size = Pt(12)
        h3_font.bold = True
    
    # User message style
    if 'UserMessage' not in styles:
        user_style = styles.add_style('UserMessage', WD_STYLE_TYPE.PARAGRAPH)
        user_style.base_style = styles['Normal']
        user_para_format = user_style.paragraph_format
        user_para_format.left_indent =Pt(0.5)
    
    # Assistant message style
    if 'AssistantMessage' not in styles:
        assistant_style = styles.add_style('AssistantMessage', WD_STYLE_TYPE.PARAGRAPH)
        assistant_style.base_style = styles['Normal']
    
    # Source style
    if 'Source' not in styles:
        source_style = styles.add_style('Source', WD_STYLE_TYPE.PARAGRAPH)
        source_style.base_style = styles['Normal']
        source_font = source_style.font
        source_font.size = Pt(9)
        source_font.color.rgb = None  # Auto color
        source_para_format = source_style.paragraph_format
        source_para_format.left_indent =Pt(0.5)
        source_para_format.space_after = Pt(0)
    if 'Preview' not in styles:
        preview_style = styles.add_style('Preview', WD_STYLE_TYPE.PARAGRAPH)
        preview_style.base_style = styles['Normal']
        preview_font = preview_style.font
        preview_font.size = Pt(8)
        preview_font.color.rgb = RGBColor(0, 102, 204)
    
    # Build chat content
    title = chat_data.get('title', 'Chat')
    messages = chat_data.get('messages', [])
    sources = chat_data.get('sources', [])
    
    # Add title
    doc.add_heading(title, level=1)
    
    # Add date/time
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph()
    
    # Add messages
    for msg in messages:
        role = msg.get('role', '')
        content = msg.get('content', '')
        
        if role == 'user':
            # User message
            p = doc.add_paragraph(f"User: ", style='UserMessage')
            _parse_inline_markdown_docx(p, content)
        elif role == 'assistant':
            # Assistant message - parse markdown
            lines = content.split('\n')
            i = 0
            while i < len(lines):
                line = lines[i]
                
                # Headers
                if line.startswith('# '):
                    text = re.sub(r'^#\s+(.+)$', r'\1', line)
                    doc.add_heading(text, level=1)
                    i += 1
                    continue
                
                if line.startswith('## '):
                    text = re.sub(r'^##\s+(.+)$', r'\1', line)
                    doc.add_heading(text, level=2)
                    i += 1
                    continue
                
                if line.startswith('### '):
                    text = re.sub(r'^###\s+(.+)$', r'\1', line)
                    doc.add_heading(text, level=3)
                    i += 1
                    continue
                
                # Blockquotes
                if line.startswith('> '):
                    text = re.sub(r'^>\s+(.+)$', r'\1', line)
                    p = doc.add_paragraph(style='AssistantMessage')
                    _parse_inline_markdown_docx(p, text)
                    p.paragraph_format.left_indent =Pt(0.3)
                    i += 1
                    continue
                
                # Unordered lists (bullets)
                if re.match(r'^\s*\*\s+', line) or re.match(r'^\s*-\s+', line):
                    while i < len(lines) and (re.match(r'^\s*\*\s+', lines[i]) or re.match(r'^\s*-\s+', lines[i])):
                        list_item_text = re.sub(r'^\s*[\*-]\s+', '', lines[i])
                        p = doc.add_paragraph(style='List Bullet')
                        _parse_inline_markdown_docx(p, list_item_text)
                        i += 1
                    continue
                
                # Ordered lists (numbers)
                if re.match(r'^\s*\d+\.\s+', line):
                    while i < len(lines) and re.match(r'^\s*\d+\.\s+', lines[i]):
                        list_item_text = re.sub(r'^\s*\d+\.\s+', '', lines[i])
                        p = doc.add_paragraph(style='List Number')
                        _parse_inline_markdown_docx(p, list_item_text)
                        i += 1
                    continue
                
                # Tables (| syntax)
                if re.match(r'^\s*\|.*\|\s*$', line):
                    result = _parse_markdown_table_docx(doc, lines[i:])
                    if result:
                        table, row_count = result
                        i += row_count
                    else:
                        p = doc.add_paragraph(style='AssistantMessage')
                        _parse_inline_markdown_docx(p, line)
                        i += 1
                    continue
                
                # Regular paragraphs
                if line.strip():
                    p = doc.add_paragraph(style='AssistantMessage')
                    _parse_inline_markdown_docx(p, line)
                
                i += 1
        
        doc.add_paragraph()
    
    # Add sources section if available
    if sources:
        doc.add_heading('Sources', level=2)
        
        for i, src in enumerate(sources, 1):
            p = doc.add_paragraph(style='Source')
            
            authors = ', '.join(src.get('authors', [])) or 'Unknown'
            title = src.get('title', 'Untitled')
            date = src.get('date', '')
            
            source_text = f"[{i}] {title} -- {authors}"
            if date:
                source_text += f" ({date})"
            
            _parse_inline_markdown_docx(p, source_text)
            
            if src.get('page_start'):
                p_page = doc.add_paragraph()
                p_page.add_run(f"  Pages: {src['page_start']}")
                if src.get('page_end') and src['page_end'] != src['page_start']:
                    p_page.add_run(f"-{src['page_end']}")
                p_page.paragraph_format.space_after = Pt(0)
            
            if src.get('item_type'):
                p_type = doc.add_paragraph()
                p_type.add_run(f"  Type: {src['item_type']}")
                p_type.paragraph_format.space_after = Pt(0)
            
            if src.get('archive'):
                p_archive = doc.add_paragraph()
                p_archive.add_run(f"  Archive: {src['archive']}")
                p_archive.paragraph_format.space_after = Pt(0)
            
            if src.get('text'):
                preview_text = _sanitize_text_for_export(src['text']).replace('\n\n', '\n')
                if preview_text:
                    p_preview = doc.add_paragraph()
                    p_preview.add_run("  Preview: ")
                    run = p_preview.add_run(preview_text)
                    run.font.color.rgb = RGBColor(0, 102, 204)
                    run.font.size = Pt(8)
                    p_preview.paragraph_format.space_after = Pt(0)
            
            if src.get('zotero_url'):
                zotero_path = src['zotero_url']
                if not zotero_path.startswith('http://') and not zotero_path.startswith('https://'):
                    full_zotero_url = f"{BASE_URL}{zotero_path}"
                else:
                    full_zotero_url = zotero_path
                p_zotero = doc.add_paragraph()
                p_zotero.add_run("  Open in Zotero: ")
                _add_hyperlink(p_zotero, full_zotero_url, full_zotero_url)
                #p_zotero.paragraph_format.space_after = Pt(0)
    
    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output


def _stream_anthropic(messages, system_prompt):
    """Stream response from Anthropic Claude."""
    import anthropic
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    with client.messages.stream(
        model=ANTHROPIC_MODEL,
        max_tokens=4096,
        system=system_prompt,
        messages=messages,
    ) as stream:
        for text in stream.text_stream:
            yield text


def _stream_openai(messages, system_prompt):
    """Stream response from OpenAI."""
    from openai import OpenAI
    client = OpenAI(api_key=OPENAI_API_KEY)
    full_messages = [{"role": "system", "content": system_prompt}] + messages
    stream = client.chat.completions.create(
        model=OPENAI_CHAT_MODEL,
        messages=full_messages,
        max_tokens=4096,
        stream=True,
    )
    for chunk in stream:
        delta = chunk.choices[0].delta
        if delta.content:
            yield delta.content


def _stream_ollama(messages, system_prompt):
    """Stream response from Ollama (local)."""
    import urllib.request
    full_messages = [{"role": "system", "content": system_prompt}] + messages
    payload = json.dumps({
        "model": OLLAMA_CHAT_MODEL,
        "messages": full_messages,
        "stream": True,
    }).encode()
    req = urllib.request.Request(
        f"{OLLAMA_BASE_URL}/api/chat",
        data=payload,
        headers={"Content-Type": "application/json"},
    )
    with urllib.request.urlopen(req, timeout=300) as resp:
        for line in resp:
            if line.strip():
                data = json.loads(line)
                content = data.get("message", {}).get("content", "")
                if content:
                    yield content


def _get_llm_stream(messages, system_prompt):
    """Get the appropriate LLM stream based on config."""
    provider = LLM_PROVIDER.lower()
    if provider == "anthropic":
        return _stream_anthropic(messages, system_prompt)
    elif provider == "ollama":
        return _stream_ollama(messages, system_prompt)
    else:
        return _stream_openai(messages, system_prompt)
 
 
@app.get("/")
@app.get("/index.html")
async def index():
    return FileResponse("static/index.html")


@app.get("/login.html")
async def login_page():
    return FileResponse("static/login.html")


@app.get("/register.html")
async def register_page():
    return FileResponse("static/register.html")


@app.get("/admin.html")
async def admin_page():
    return FileResponse("static/admin.html")


@app.post("/api/login")
async def login(form_data: OAuth2PasswordRequestForm = Depends()):
    """Login and get JWT token."""
    user = authenticate_user(form_data.username, form_data.password)
    if not user:
        raise HTTPException(
            status_code=401,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token = create_access_token(data={"sub": user["username"]})
    return {
        "access_token": access_token,
        "token_type": "bearer",
        "username": user["username"],
        "is_admin": user.get("is_admin", False)
    }


@app.post("/api/register")
async def register(request: RegisterRequest):
    """Register a new user account (requires admin approval)."""
    existing_user = get_auth_user(request.username)
    if existing_user:
        raise HTTPException(
            status_code=400,
            detail="Username already registered"
        )
    user = create_user(request.username, request.password, request.email)
    if not user:
        raise HTTPException(
            status_code=500,
            detail="Failed to create user"
        )
    return {
        "message": "Registration submitted for admin approval",
        "username": request.username,
        "email": request.email
    }


@app.get("/api/logout")
async def logout(current_user = Depends(get_current_user)):
    """Logout endpoint."""
    return {"message": "Logged out successfully"}


@app.get("/api/me")
async def get_me(current_user = Depends(get_current_user)):
    """Get current user info."""
    return {
        "username": current_user["username"],
        "email": current_user.get("email", ""),
        "is_admin": current_user.get("is_admin", False),
        "is_approved": current_user.get("is_approved", True)
    }


def require_admin(current_user = Depends(get_current_user)):
    """Dependency that requires the user to be an admin."""
    if not current_user.get("is_admin", False):
        raise HTTPException(status_code=403, detail="Admin access required")
    return current_user


@app.get("/api/admin/users")
async def list_users(admin = Depends(require_admin)):
    """List all users (admin only)."""
    users = get_all_users()
    return {
        "users": [
            {
                "username": u["username"],
                "email": u.get("email", ""),
                "is_admin": u.get("is_admin", False),
                "is_approved": u.get("is_approved", True),
                "created_at": u.get("created_at", "")
            }
            for u in users
        ]
    }


@app.post("/api/admin/users/{username}/approve")
async def approve_user_endpoint(username: str, admin = Depends(require_admin)):
    """Approve a pending user (admin only)."""
    if approve_user(username):
        return {"message": f"User {username} approved"}
    raise HTTPException(status_code=404, detail="User not found")


@app.delete("/api/admin/users/{username}")
async def delete_user_endpoint(username: str, admin = Depends(require_admin)):
    """Delete a user (admin only)."""
    if delete_user(username):
        return {"message": f"User {username} deleted"}
    raise HTTPException(status_code=404, detail="User not found")


@app.get("/api/chats")
async def list_chats(current_user = Depends(get_current_user)):
    """List all chats for the current user."""
    chats = get_user_chats(current_user["username"])
    return {"chats": chats}


@app.get("/api/chats/{chat_id}")
async def get_chat_endpoint(chat_id: str, current_user = Depends(get_current_user)):
    """Get a specific chat with messages."""
    chat = get_chat(chat_id, current_user["username"])
    if not chat:
        raise HTTPException(status_code=404, detail="Chat not found")
    return chat


class SyncChatRequest(BaseModel):
    title: str
    messages: list
    sources: list


@app.post("/api/chats/{chat_id}/sync")
async def sync_chat_endpoint(chat_id: str, request: SyncChatRequest, current_user = Depends(get_current_user)):
    """Sync a chat (create or update with all messages)."""
    sync_chat(
        chat_id,
        current_user["username"],
        request.title,
        request.messages,
        request.sources
    )
    return {"message": "Chat synced"}


@app.delete("/api/chats/{chat_id}")
async def delete_chat_endpoint(chat_id: str, current_user = Depends(get_current_user)):
    """Delete a chat."""
    if delete_chat(chat_id, current_user["username"]):
        return {"message": "Chat deleted"}
    raise HTTPException(status_code=404, detail="Chat not found")


@app.get("/api/filters")
async def filters(current_user = Depends(get_current_user)):
    """Return available filter values for sidebar dropdowns."""
    aliases = get_archive_aliases()

    archive_options = []
    for acronym, full_name in sorted(aliases.items(), key=lambda x: x[1]):
        archive_options.append({
            'value': acronym,
            'label': f"{acronym.upper()} -- {full_name}",
        })

    item_types = [
        'book', 'bookSection', 'conferencePaper', 'document',
        'hearing', 'journalArticle', 'letter', 'manuscript',
        'newspaperArticle', 'report', 'statute', 'thesis', 'webpage',
    ]

    return JSONResponse({
        'archives': archive_options,
        'item_types': item_types,
    })


@app.post("/api/chat")
async def chat(request: Request, current_user = Depends(get_current_user)):
    """Search + stream LLM response via SSE."""
    body = await request.json()
    message = body.get('message', '').strip()
    prev_conversation = body.get('conversation', [])
    existing_sources = body.get('sources', [])
    filter_vals = body.get('filters', {})
    top_k = body.get('top_k', 10)

    if not message:
        return JSONResponse({'error': 'Empty message'}, status_code=400)

    async def generate():
        try:
            results = run_search(
                message,
                top_k=top_k,
                item_type=filter_vals.get('item_type') or None,
                author=filter_vals.get('author') or None,
                tag=filter_vals.get('tag') or None,
                collection=filter_vals.get('collection') or None,
                archive=filter_vals.get('archive') or None,
                date_from=filter_vals.get('date_from') or None,
                date_to=filter_vals.get('date_to') or None,
            )

            # Build lookup set of existing sources to deduplicate
            existing_source_keys = {
                f"{s.get('zotero_key', '')}_{s.get('chunk_index', 0)}" 
                for s in existing_sources
            }
            
            # Assign source_num to new sources based on offset
            source_offset = len(existing_sources)
            client_sources = []
            for r in results:
                formatted = _format_source_for_client(r)
                compound_key = f"{formatted['zotero_key']}_{formatted['chunk_index']}"
                if compound_key in existing_source_keys:
                    # Source already exists, use its source_num
                    existing = next((s for s in existing_sources if f"{s.get('zotero_key', '')}_{s.get('chunk_index', 0)}" == compound_key), None)
                    if existing and 'source_num' in existing:
                        formatted['source_num'] = existing['source_num']
                    else:
                        # Fallback: increment offset
                        source_offset += 1
                        formatted['source_num'] = source_offset
                else:
                    # New unique source, assign next number
                    source_offset += 1
                    formatted['source_num'] = source_offset
                client_sources.append(formatted)
            
            yield f"data: {json.dumps({'type': 'sources', 'sources': client_sources})}\n\n"

            source_context = _build_source_context(client_sources)

            messages = []
            for msg in prev_conversation:
                messages.append({
                    'role': msg['role'],
                    'content': msg['content'],
                })

            user_content = f"{source_context}\n\nUser question: {message}"
            messages.append({'role': 'user', 'content': user_content})

            for text in _get_llm_stream(messages, SYSTEM_PROMPT):
                yield f"data: {json.dumps({'type': 'delta', 'text': text})}\n\n"

            yield f"data: {json.dumps({'type': 'done'})}\n\n"

        except Exception as e:
            logger.exception("Chat error")
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")




@app.get("/api/chat/download")
async def download_chat(
    chat_id: str = None,
    format: str = 'odt',
    current_user = Depends(get_current_user)
):
    """Download current chat as ODT or DOCX file."""
    if not chat_id:
        raise HTTPException(status_code=400, detail="chat_id parameter required")
    
    chat_data = get_chat(chat_id, current_user["username"])
    if not chat_data:
        raise HTTPException(status_code=404, detail="Chat not found")
    
    format = format.lower()
    
    if format == 'docx':
        docx_content = _chat_to_docx(chat_data)
        
        return Response(
            content=docx_content.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=chat_{chat_id[:10]}.docx"
            }
        )
    else:
        odt_content = _chat_to_odt(chat_data)
        
        return Response(
            content=odt_content.getvalue(),
            media_type="application/vnd.oasis.opendocument.text",
            headers={
                "Content-Disposition": f"attachment; filename=chat_{chat_id[:10]}.odt"
            }
        )


@app.get("/api/chat/download/docx")
async def download_chat_docx(
    chat_id: str = None,
    current_user = Depends(get_current_user)
):
    """Download current chat as DOCX file (shortcut endpoint)."""
    if not chat_id:
        raise HTTPException(status_code=400, detail="chat_id parameter required")
    
    chat_data = get_chat(chat_id, current_user["username"])
    if not chat_data:
        raise HTTPException(status_code=404, detail="Chat not found")
    
    docx_content = _chat_to_docx(chat_data)
    
    return Response(
        content=docx_content.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=chat_{chat_id[:10]}.docx"
        }
    )


@app.get("/zotero/pdf/{key}")
async def open_pdf(key: str, page: int = 0):
    """Open a PDF in Zotero via zotero:// URL."""
    url = f"zotero://open-pdf/library/items/{key}"
    if page:
        url += f"?page={page}"
    # macOS
    if sys.platform == "darwin":
        subprocess.Popen(['open', url])
    # Linux
    elif sys.platform.startswith("linux"):
        subprocess.Popen(['xdg-open', url])
    # Windows
    elif sys.platform == "win32":
        os.startfile(url)
    return HTMLResponse(
        '<html><body><p>Opened in Zotero.</p>'
        '<script>window.close()</script></body></html>'
    )


@app.get("/zotero/item/{key}")
async def open_item(key: str):
    """Open a Zotero item via zotero:// URL."""
    url = f"zotero://select/library/items/{key}"
    if sys.platform == "darwin":
        subprocess.Popen(['open', url])
    elif sys.platform.startswith("linux"):
        subprocess.Popen(['xdg-open', url])
    elif sys.platform == "win32":
        os.startfile(url)
    return HTMLResponse(
        '<html><body><p>Opened in Zotero.</p>'
        '<script>window.close()</script></body></html>'
    )


if __name__ == '__main__':
    import uvicorn

    print("Initializing search pipeline...", end=' ', flush=True)
    init_pipeline()
    print("done.")
    print(f"Using LLM provider: {LLM_PROVIDER}")
    print("Open http://localhost:5001 in your browser.")
    baseurl = BASE_URL.split(":")[1].split("//")[1] if ":" in BASE_URL else "127.0.0.1"
    port = int(BASE_URL.split(":")[-1]) if ":" in BASE_URL else 5001
    uvicorn.run(app, host=baseurl, port=port, log_level="info")
